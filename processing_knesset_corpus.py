from docx import Document
import zipfile
import os
import pandas
import re
import json
import sys

def get_docx(folder_path):
    try:

        info =[]
        # go throw all the documants in the directory
        current_path =os.path.join(os.getcwd(),folder_path)
        
        for file_number,filename in enumerate(os.listdir(current_path)):
            if filename.endswith('docx'):
                
                attributes = filename.split('_')# get the attributes of the file
                #then assigen it to the right variable
                if attributes[1] == 'ptv':  # committee or plenary
                        type = 'committee'
                elif attributes[1] == 'ptm':
                        type = 'plenary'
                else:
                    type = '-1'

                text = Document(os.path.join(current_path,filename))
                
                info_item = {
                        'file_name': filename,
                        'knesset number': int(attributes[0]),
                        'type': type,
                        'text': text,
                        'file_number': attributes[2].replace('.docx', '')
                    }
                info.append(info_item)
        return info
    except Exception as e:
        print(f'Exception in get_all_docx_in_current_foleder: {e}')

# Function to the number after we find "הישיבה ה" or "'פרוטוקול מס"
def get_next_word(text, position):
    # Find the start of the next word
    word_start = position


    while word_start < len(text) and text[word_start].isspace(): # Skip spaces
        word_start += 1

    # If reached the end of the text, return empty
    if word_start >= len(text):
        return "-1"

    # Find the end of the next word
    word_end = word_start
    while word_end < len(text) and not text[word_end].isspace():
        word_end += 1

    # Return the next continuous word
    return text[word_start:word_end]

def is_underlined(par):
    try:
        # Check if all text runs are underlined
        all_underlined = True  # Assume all are underlined unless proven otherwise
        for run in par.runs:
            if not run.underline:
                all_underlined = False  # If any run is not underlined, then it's not fully underlined

        if all_underlined:
            return True

        # If any run is underlined, consider it partially underlined
        for run in par.runs:
            if run.underline:
                return True

        # Check the paragraph style for underline attributes
        current_style = par.style
        while current_style:
            if current_style.font.underline:
                return True  # If underline is defined at the style level

            # Check base styles for inherited underline
            current_style = current_style.base_style

        return False  # Default to not underlined
    except Exception as e:
        print(f"Exception in is_underlined: {e}")
        return False

def clean_name(name): # clean
    try:
        name = name.strip()
        comps = name.split(' ') #split name to words/comps
        clean_name = ""
        open_parentheses =False 

        for comp in comps:
            if comp == '':
                continue

            if '(' in comp: # This is a closing parentheses, backwads since its in hebrew
                open_parentheses = False
                continue
            if open_parentheses:
                continue
            if '"' not in comp and '”' not in comp:# if name isnt a shortcut
                if '\'' == comp[-1] and len(comp) <4: #then this means that it might be the person numbering of his possition
                    clean_name = '' # if it is then remove what we collected before
                    continue

                    #if the number is more than 2 digits in hebrew this means this code will not capture it
                    #if someone name has 2 latters and ends with ' then this code will not capture it (unlikely to happen)
                    #if position came after the name then this code will not capture the name

                if ")" in comp: # This is an opening parentheses, backwads since its in hebrew
                    if "(" in comp:# if we have closing parentheses then skip it
                        continue
                    else:# else we must take the following comps until we see closing parentheses
                        open_parentheses = True

                elif comp == "-" or comp == '–' or comp == '~' or comp ==',':#if the name has a dash then take the first part
                    break
                else:
                    clean_name += comp+" "

        clean_name = clean_name.strip()
        if ',' in clean_name:
            return ''
        if clean_name != "" and clean_name.find(':')+1 == len(clean_name): # Remove the colon if it's the last character, check if name isnt empty
            clean_name = clean_name[:-1]

        return clean_name.strip()
    except Exception as e:
        print(f'Exception in clear_name: {e}')

def split_paragrph(par):
    try:
        new_sentence = ''
        seperators = '.؟!?!' # use this to check start and end of sentences
        qutoed = False
        par_parts = par.text.strip().split(' ')
        sentece_list = []
        for part in par_parts:
            
            if part =='':# if empty then skip
                continue
            
            new_sentence += part +" " # Collect sentence
            if '"' == part[0] and qutoed == False: 
                qutoed = True
            if '"' == part[-1] or (len(part)>=2 and part[-2] == '"' and part[-1] in seperators):
                # if the second to last char is a " and last is a seperator, or if the last is a quote then we end the quet 
                qutoed = False

            if part[-1]  in seperators or (len(part)>=2 and part[-2] in seperators): # If we reached the end of the sentence, save it
                if qutoed == False: # if were still in quotes, we dont save yet
                    sentece_list.append(new_sentence.strip())
                    new_sentence = ''

        #if we start a quote but it didnt end, save the text
        if qutoed:
            sentece_list.append(new_sentence)
        return sentece_list
    except Exception as e:
        print(f'exception in split_paragrph: {e}')

def remove_tags(text, tags):
    # Strip leading/trailing spaces
    cleaned_text = text.strip()
    for tag in tags:
        # Remove specific tag from start
        if cleaned_text.startswith(tag):
            cleaned_text = cleaned_text[len(tag):].strip()  # Remove the tag and strip spaces
        
        # Remove specific tag from end
        if cleaned_text.endswith(tag):
            cleaned_text = cleaned_text[:-len(tag)].strip()  # Remove the tag and strip spaces
        
    return cleaned_text


def clean_text(txt):
    try:
        if txt == '':
            return ''
        allowed = re.compile('[א-ת0-9!"#$%&\'()*+,-./:;<=>?@[\\]_`{|}~– ]+') # Allowed characters
        occurences = re.findall(allowed, txt) # Find all allowed characters
        if len(occurences) != 1: # If there are more than one occurence that means we have unwated characters in the text or in between the text
            return ''
        
        filtered_txt = occurences[0] # Get the first and only occurence, we must check if its actually hebrew or not
        heb_txt = False # Check if the text is in hebrew
        heb_letters = [chr(code) for code in range(0x05D0, 0x05EA + 1)]

        for letter in filtered_txt:
            if letter in heb_letters: 
                heb_txt = True # This means that the text is in hebrew since theres one occurence
                break
        if heb_txt == False:
            return ''
        
        # Check for special cases
        cases = ['- - -','- -' , '– – –','– –' ,'– – –','– –' ]
        if any(case in filtered_txt for case in cases): # If the text contains any of the special cases, return empty
            return ''
        return txt


    except Exception as e:
        print(f'Exception in clean_text {e}')

if __name__ == "__main__":
    try:
        #print(sys.argv)
        if len(sys.argv) !=3:
            print('Incorrect input, please enter the folder path and the output path.')
            sys.exit(1)

        # Add Check if folder is valid

       # CHECK IF ITS IN CORRECT JSONL FORMAT
        folder_path = sys.argv[1]
        output_path = sys.argv[2]
        info = get_docx(folder_path)
        cpy = info.copy()

        tags = ["<< דובר >>", "<< נושא >>", "<< יור >>", "<< דובר_המשך >>", "<< אורח >>"] # Tags to remove from the text
        common_pos = ["סדר-היום", "סדר היום", "נכחו", "חברי", "מנהל", "רישום", "משתתפים", "מוזמנים", "ייעוץ", "יועץ", "קצרנית", "יועצת", "קצרן"] # List of all the common positions
        target_words = ["הישיבה ה", "פרוטוקול מס'"]  # Search for the protocal number
        jsonl_data = []
        names = []

        #clean_text('  שלום  ')
        #clean_text('  hi  ')
        #clean_text('  hi כן  לט')
        #clean_text('  ללאלhi  ')
        #clean_text('לא - -')
        #clean_text('שדג - - -')
        #clean_text( ' hi  לא  hi ')
        #clean_text(' לא  hi  לא  ')

        for doc_num, doc in enumerate(info):
            knesset_number = doc['knesset number']
            protocol_type = doc['type']
            file_number = doc['file_number']
            protocol_name = doc['file_name']
            protocol_number = -1 # Default value
            
            speakers_order = [] # list to hold the order of the speakers
    
            for par in doc['text'].paragraphs:
                text = par.text.strip() # Remove leading and trailing spaces

                if text.startswith('<') or text.startswith('>'): # Sometimes the text starts with < and ends with >, its probably caused by the conversion from doc to docx so we remove it
                    text = text[1:-1] 

                if target_words[0] in text:
                    position = text.find(target_words[0])
                    next_word = get_next_word(text, position + len(target_words[0]))
                    print(f"Found in doc {doc_num}: {text}, NEXT {next_word}")
                    protocol_number = next_word
                    break
                
                if target_words[1] in text:
                    position = text.find(target_words[1])
                    next_word = get_next_word(text, position + len(target_words[1]))
                    print(f"Found in doc {doc_num}: {text}, NEXT {next_word}")
                    protocol_number = next_word
                    break
            
            

            # Extract speakers and text

            prev_speaker ='' # name of the current_speaker
            speaker_text = {} # dict to hold the text of each speaker
            
            #if doc['file_number'] != '3841247':
            #    continue
            #for i in range(100):

            ########## REMOVE THIS ##########
            ########## REMOVE THIS ##########
            ########## REMOVE THIS ##########
            names.append({'docx':doc['file_number']}) ########## DONT FORGE TTO REMOVE 
            ########## REMOVE THIS ##########
            ########## REMOVE THIS ##########
            ########## REMOVE THIS ##########


            for par in doc['text'].paragraphs:
                text = par.text
                text = remove_tags(text, tags)
                #text = doc['text'].paragraphs[i].text
                if text.startswith('<') or text.startswith('>'): # Sometimes the text starts with < and ends with >, its probably caused by the conversion from doc to docx so we remove it
                    text = text[1:-1] 
                
                index = text.find(":")
                if index>=0:  # if the last char is : and the whole text is underlined then this is a speaker
                    
                    if index== len(text) -1 and is_underlined(par):
                        if any(pos in text for pos in common_pos): # if the text contains any of the common positions then skip
                            #names.append({'common_pos':text})
                            continue
                        new_name = clean_name(text)
                        #names.append({'names':text,
                        #              'clear_name':new_name})
                        
                        if new_name != '':
                            prev_speaker = new_name
                        
                        else: # This happens when the name is in the setnence (has ,)
                            split_txt = split_paragrph(par)
                            #speaker_text[prev_speaker].append(split_txt)

                            if prev_speaker != '':
                                for sent in split_txt:
                                    filtered = clean_text(sent)
                                    if filtered != '':
                                        speaker_text[prev_speaker].append(filtered)

                        if prev_speaker not in list(speaker_text.keys()): # if the speaker is not in the dict then add him
                            speaker_text[prev_speaker] = []  

                        if prev_speaker not in speakers_order:
                            speakers_order.append(prev_speaker)

                    elif prev_speaker != '': # if the speaker is not empty then this is a continuation of his speech
                        split_txt = split_paragrph(par)
                        for sent in split_txt:
                            filtered = clean_text(sent)
                            if filtered != '':
                                speaker_text[prev_speaker].append(filtered)

                elif prev_speaker != '':# if we have a speaker then add the text to his name
                    split_txt = split_paragrph(par)
                    for sent in split_txt:
                        filtered = clean_text(sent)
                        if filtered != '':
                            speaker_text[prev_speaker].append(filtered)
            #info[doc_num]['speaker_data'] = speaker_text# save the data


            # Put this in the correct place
            # Append the data to the jsonl_data list

            for speaker in speakers_order:
                for text in speaker_text[speaker]:
                    jsonl_data.append({
                        'protocol_name': protocol_name,
                        'knesset_number': knesset_number,
                        'protocol_type': protocol_type,
                        'protocol_number': protocol_number,
                        'speaker_name': speaker,
                        'sentence_text': text
                    })
        
        with open(output_path, 'w', encoding='utf-8') as jsonl_file:
            for data_item in jsonl_data: # change back
                # Convert the dictionary to a JSON-formatted string
                json_line = json.dumps(data_item, ensure_ascii=False)
        
                # Write the JSON string to the file with a newline separator
                jsonl_file.write(json_line + '\n')

    except Exception as e:
        # Handle any exception
        print(f"An error occurred in main: {e}")
